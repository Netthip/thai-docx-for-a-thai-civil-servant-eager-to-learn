#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
thai_official.py — สร้าง "หนังสือราชการไทย" ตามรูปแบบระเบียบงานสารบรรณ
ต่อยอดจาก thai_docx.py (ใช้ helper เดิม: ฟอนต์ครบฝั่ง Complex Script + ข้อความเต็มบรรทัด)

รองรับ 2 ชนิดที่ใช้บ่อยที่สุด
  • create_memo()   — บันทึกข้อความ (หนังสือภายใน)
  • create_letter() — หนังสือภายนอก (มีคำขึ้นต้น/คำลงท้าย)

จุดเด่น
  - วางโครงตามแบบฟอร์มจริง: ส่วนราชการ / ที่ / วันที่ / เรื่อง / เรียน / (อ้างถึง) /
    (สิ่งที่ส่งมาด้วย) / เนื้อหา / คำลงท้าย / ลงชื่อ
  - ภาษาไทยเต็มบรรทัด ฟอนต์ถูกต้องทุกฝั่ง (ยืมจาก thai_docx.py)
  - แปลงเลขอารบิกเป็นเลขไทย (๐–๙) และจัดวันที่แบบไทย (พ.ศ.) ได้
  - ตราครุฑใส่เองได้ (ระบุ path รูป) — ถ้าไม่ใส่ จะเว้นที่ไว้ตามแบบฟอร์ม

หมายเหตุเรื่องตราครุฑ: รีโปนี้ไม่ฝังรูปครุฑมาให้ (เป็นตราแผ่นดิน ใช้ตามระเบียบ)
ผู้ใช้ระบุไฟล์รูปของหน่วยงานเองผ่านพารามิเตอร์ garuda_image

เผยแพร่เป็นสาธารณะ (MIT License)
"""

from __future__ import annotations

import os
import sys
import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

# ใช้ helper จาก thai_docx.py ที่อยู่โฟลเดอร์เดียวกัน
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from thai_docx import (  # noqa: E402
    set_run_font,
    insert_zwsp,
    resolve_font,
    _apply_page_setup,
    _set_styles,
    _safe_stdout,
)

# ─────────────────────────────────────────────────────────────────────────────
# เลขไทย / วันที่ไทย
# ─────────────────────────────────────────────────────────────────────────────
_THAI_DIGIT_MAP = str.maketrans("0123456789", "๐๑๒๓๔๕๖๗๘๙")
_THAI_MONTHS = [
    "", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
    "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม",
]


def to_thai_digits(text) -> str:
    """แปลงเลขอารบิก 0-9 ในข้อความเป็นเลขไทย ๐-๙ (ส่วนอื่นคงเดิม)"""
    return str(text).translate(_THAI_DIGIT_MAP)


def thai_date(d: datetime.date = None, thai_digits: bool = True) -> str:
    """จัดรูปวันที่แบบราชการ: '๑ กรกฎาคม ๒๕๖๙' (ปี พ.ศ.)

    d=None จะใช้วันที่วันนี้ ส่ง datetime.date เองได้เพื่อผลลัพธ์คงที่ (เช่นในเทสต์)
    """
    if d is None:
        d = datetime.date.today()
    s = f"{d.day} {_THAI_MONTHS[d.month]} {d.year + 543}"
    return to_thai_digits(s) if thai_digits else s


# ─────────────────────────────────────────────────────────────────────────────
# ค่าตั้งต้นการจัดหน้าหนังสือราชการ (อิงแนวทางระเบียบงานสารบรรณ)
#   TH Sarabun New 16pt, A4, ขอบ บน 2.5 / ล่าง 2.0 / ซ้าย 3.0 / ขวา 2.0 ซม.
# ─────────────────────────────────────────────────────────────────────────────
_OFFICIAL_DEFAULTS = dict(
    font_name="TH Sarabun New",
    font_size=16,
    page_size="A4",
    margins=dict(top=2.5, bottom=2.0, left=3.0, right=2.0),
)


def _normalize_body(body) -> list:
    """รับเนื้อหาได้หลายแบบ → คืน list ของย่อหน้า (str)
    - str: แยกย่อหน้าด้วยบรรทัดว่าง
    - list[str]: ถือว่าแต่ละตัวคือหนึ่งย่อหน้า
    """
    if body is None:
        return []
    if isinstance(body, str):
        blocks = [b.strip() for b in body.replace("\r\n", "\n").split("\n\n")]
        return [b for b in blocks if b]
    return [str(b) for b in body]


def _content_width_cm(cfg) -> float:
    """ความกว้างพื้นที่พิมพ์ (ซม.) = กว้างกระดาษ - ขอบซ้าย - ขอบขวา"""
    page_w = 21.0 if cfg["page_size"] == "A4" else 21.59
    m = cfg["margins"]
    return page_w - m["left"] - m["right"]


def _new_doc(cfg):
    font_name = resolve_font(cfg["font_name"])
    font_size = cfg["font_size"]
    doc = Document()
    _set_styles(doc, font_name, font_size)
    _apply_page_setup(doc.sections[0], cfg["page_size"], cfg["margins"])
    return doc, font_name, font_size


def _blank(doc, size_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if size_pt:
        for r in [p.add_run("")]:
            set_run_font(r, size=size_pt)
    return p


def _label_value(doc, label, value, font_name, font_size, *, gap_runs=2, bold_label=True):
    """ย่อหน้ารูปแบบ 'ป้ายตัวหนา  ค่า' เช่น 'เรื่อง  ขออนุมัติ...' """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    set_run_font(p.add_run(label), name=font_name, size=font_size, bold=bold_label or None)
    set_run_font(p.add_run(" " * gap_runs), name=font_name, size=font_size)
    set_run_font(p.add_run(insert_zwsp(str(value))), name=font_name, size=font_size)
    return p


def _hr(doc):
    """เส้นคั่นแนวนอน (ใช้ใต้บรรทัด 'เรื่อง' ในบันทึกข้อความ)"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def _add_garuda(doc, garuda_image, height_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    if garuda_image and os.path.exists(garuda_image):
        p.add_run().add_picture(garuda_image, height=Cm(height_cm))
    else:
        # ไม่มีรูป — เว้นที่ไว้ตามแบบฟอร์ม พร้อมหมายเหตุจาง ๆ
        set_run_font(
            p.add_run(f"[ ตราครุฑ สูง {to_thai_digits(height_cm)} ซม. ]"),
            size=12, italic=True, color=RGBColor(0xAA, 0xAA, 0xAA),
        )
    return p


def _add_body(doc, body, font_name, font_size, *, first_line_indent_cm, line_spacing):
    for para in _normalize_body(body):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = line_spacing
        pf.space_after = Pt(0)
        if first_line_indent_cm:
            pf.first_line_indent = Cm(first_line_indent_cm)
        set_run_font(p.add_run(insert_zwsp(para)), name=font_name, size=font_size)


def _add_signature(doc, font_name, font_size, *, name, position, position2,
                   closing=None, block_left_cm):
    """บล็อกลงนาม จัดให้อยู่ค่อนไปทางขวา (เยื้องซ้าย block_left_cm)"""
    def line(text, *, indent_extra=0.0):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(block_left_cm + indent_extra)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_run_font(p.add_run(insert_zwsp(text)), name=font_name, size=font_size)
        return p

    if closing:                      # คำลงท้าย เช่น "ขอแสดงความนับถือ"
        line(closing, indent_extra=0.8)
        _blank(doc, 0)
    _blank(doc, 0)                   # เว้นที่เซ็นชื่อ
    if name:
        line(f"({name})", indent_extra=1.0)
    if position:
        line(position)
    if position2:
        line(position2)


# ─────────────────────────────────────────────────────────────────────────────
# 1) บันทึกข้อความ
# ─────────────────────────────────────────────────────────────────────────────
def create_memo(
    output_path: str,
    *,
    department: str,
    subject: str,
    to: str,
    body,
    doc_number: str = "",
    date=None,
    signer_name: str = "",
    signer_position: str = "",
    signer_position2: str = "",
    garuda_image: str = None,
    font_name: str = None,
    font_size: int = None,
    thai_digits: bool = False,
):
    """
    สร้าง "บันทึกข้อความ" (.docx)

    department      ส่วนราชการ (เช่น 'กองคลัง สำนัก...')
    subject         เรื่อง
    to              เรียน (เช่น 'ผู้อำนวยการสำนัก...')
    body            เนื้อหา — str (แยกย่อหน้าด้วยบรรทัดว่าง) หรือ list[str]
    doc_number      เลขที่หนังสือ (ช่อง 'ที่')
    date            วันที่ — str ที่จัดรูปแล้ว, datetime.date, หรือ None=วันนี้
    signer_name     ชื่อผู้ลงนาม (จะใส่ในวงเล็บ)
    signer_position ตำแหน่ง (รองรับ 2 บรรทัดผ่าน signer_position2)
    garuda_image    path รูปตราครุฑ (ไม่ใส่ก็ได้)
    thai_digits     True = แปลงเลขใน'เลขที่/วันที่' เป็นเลขไทย
    """
    cfg = dict(_OFFICIAL_DEFAULTS)
    if font_name is not None:
        cfg["font_name"] = font_name
    if font_size is not None:
        cfg["font_size"] = font_size

    doc, fname, fsize = _new_doc(cfg)
    content_w = _content_width_cm(cfg)

    # ── หัวเอกสาร: ครุฑ + คำว่า "บันทึกข้อความ" ──
    _add_garuda(doc, garuda_image, height_cm=1.5)
    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head.paragraph_format.space_after = Pt(6)
    set_run_font(head.add_run("บันทึกข้อความ"), name=fname, size=fsize + 13, bold=True)

    # ── ส่วนหัวข้อมูล ──
    if thai_digits:
        doc_number = to_thai_digits(doc_number)
    date_str = date if isinstance(date, str) else thai_date(date, thai_digits=thai_digits)

    _label_value(doc, "ส่วนราชการ", department, fname, fsize)

    # บรรทัด "ที่ ... วันที่ ..." สองคอลัมน์ด้วย tab ขวา
    row = doc.add_paragraph()
    row.paragraph_format.space_after = Pt(0)
    row.paragraph_format.line_spacing = 1.0
    row.paragraph_format.tab_stops.add_tab_stop(Cm(content_w), WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(row.add_run("ที่"), name=fname, size=fsize, bold=True)
    set_run_font(row.add_run("  " + insert_zwsp(doc_number or "")), name=fname, size=fsize)
    set_run_font(row.add_run("\t"), name=fname, size=fsize)
    set_run_font(row.add_run("วันที่"), name=fname, size=fsize, bold=True)
    set_run_font(row.add_run("  " + insert_zwsp(date_str)), name=fname, size=fsize)

    _label_value(doc, "เรื่อง", subject, fname, fsize)
    _hr(doc)

    # ── เรียน ──
    _blank(doc, 0)
    _label_value(doc, "เรียน", to, fname, fsize)
    _blank(doc, 0)

    # ── เนื้อหา ──
    _add_body(doc, body, fname, fsize, first_line_indent_cm=1.25, line_spacing=1.0)

    # ── ลงชื่อ ──
    _blank(doc, 6)
    _add_signature(
        doc, fname, fsize,
        name=signer_name, position=signer_position, position2=signer_position2,
        closing=None, block_left_cm=content_w * 0.52,
    )

    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# 2) หนังสือภายนอก
# ─────────────────────────────────────────────────────────────────────────────
def create_letter(
    output_path: str,
    *,
    subject: str,
    to: str,
    body,
    doc_number: str = "",
    sender_lines=None,
    date=None,
    reference: str = "",
    enclosure: str = "",
    salutation: str = None,
    closing: str = "ขอแสดงความนับถือ",
    signer_name: str = "",
    signer_position: str = "",
    signer_position2: str = "",
    contact_lines=None,
    garuda_image: str = None,
    font_name: str = None,
    font_size: int = None,
    thai_digits: bool = False,
):
    """
    สร้าง "หนังสือภายนอก" (.docx) ตามรูปแบบหนังสือราชการ

    subject       เรื่อง
    to            เรียน (ผู้รับ)
    body          เนื้อหา — str หรือ list[str] (ย่อหน้าแรกควรเป็นคำขึ้นต้น เช่น 'ด้วย...'/'ตามที่...')
    doc_number    เลขที่หนังสือ (ช่อง 'ที่')
    sender_lines  list[str] ชื่อ/ที่อยู่หน่วยงานผู้ส่ง (บล็อกขวาบน)
    date          วันที่ (กึ่งกลาง) — str / datetime.date / None=วันนี้
    reference     อ้างถึง (ไม่ใส่ก็ได้)
    enclosure     สิ่งที่ส่งมาด้วย (ไม่ใส่ก็ได้)
    closing       คำลงท้าย (ดีฟอลต์ 'ขอแสดงความนับถือ')
    contact_lines list[str] ส่วนล่างซ้าย เช่น ['กองคลัง', 'โทร. ๐ ๒๒๐๒ ๑๒๓๔']
    """
    cfg = dict(_OFFICIAL_DEFAULTS)
    if font_name is not None:
        cfg["font_name"] = font_name
    if font_size is not None:
        cfg["font_size"] = font_size

    doc, fname, fsize = _new_doc(cfg)
    content_w = _content_width_cm(cfg)

    if thai_digits:
        doc_number = to_thai_digits(doc_number)
    date_str = date if isinstance(date, str) else thai_date(date, thai_digits=thai_digits)

    # ── ครุฑกลางหน้า (3 ซม.) ──
    _add_garuda(doc, garuda_image, height_cm=3.0)

    # ── แถวบน: "ที่ ..." ซ้าย + บล็อกที่อยู่ผู้ส่งขวา (ตารางไร้เส้น 2 ช่อง) ──
    sender_lines = sender_lines or []
    top = doc.add_table(rows=1, cols=2)
    top.autofit = True
    left_cell, right_cell = top.cell(0, 0), top.cell(0, 1)

    lp = left_cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    set_run_font(lp.add_run("ที่"), name=fname, size=fsize, bold=True)
    set_run_font(lp.add_run("  " + insert_zwsp(doc_number or "")), name=fname, size=fsize)

    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for idx, line in enumerate(sender_lines):
        rp = right_cell.paragraphs[0] if idx == 0 else right_cell.add_paragraph()
        rp.paragraph_format.space_after = Pt(0)
        rp.paragraph_format.line_spacing = 1.0
        set_run_font(rp.add_run(insert_zwsp(line)), name=fname, size=fsize)

    # ── วันที่ กึ่งกลาง ──
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp.paragraph_format.space_before = Pt(2)
    dp.paragraph_format.space_after = Pt(2)
    set_run_font(dp.add_run(insert_zwsp(date_str)), name=fname, size=fsize)

    # ── เรื่อง / เรียน / อ้างถึง / สิ่งที่ส่งมาด้วย ──
    _label_value(doc, "เรื่อง", subject, fname, fsize)
    _label_value(doc, "เรียน", to, fname, fsize)
    if reference:
        _label_value(doc, "อ้างถึง", reference, fname, fsize)
    if enclosure:
        _label_value(doc, "สิ่งที่ส่งมาด้วย", enclosure, fname, fsize)
    if salutation:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(0)
        set_run_font(sp.add_run(insert_zwsp(salutation)), name=fname, size=fsize)
    _blank(doc, 2)

    # ── เนื้อหา ──
    _add_body(doc, body, fname, fsize, first_line_indent_cm=2.5, line_spacing=1.0)

    # ── คำลงท้าย + ลงชื่อ ──
    _blank(doc, 6)
    _add_signature(
        doc, fname, fsize,
        name=signer_name, position=signer_position, position2=signer_position2,
        closing=closing, block_left_cm=content_w * 0.52,
    )

    # ── ส่วนล่างซ้าย: หน่วยงานเจ้าของเรื่อง / โทร. ──
    if contact_lines:
        _blank(doc, 12)
        for line in contact_lines:
            cp = doc.add_paragraph()
            cp.paragraph_format.space_after = Pt(0)
            cp.paragraph_format.line_spacing = 1.0
            set_run_font(cp.add_run(insert_zwsp(line)), name=fname, size=fsize - 2)

    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# CLI สั้น ๆ (เดโม) — ปกติแนะนำเรียกผ่านโค้ด เพราะฟิลด์เยอะ
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    _safe_stdout()
    out = sys.argv[1] if len(sys.argv) > 1 else "memo_demo.docx"
    create_memo(
        out,
        department="กองคลัง สำนักงานเลขานุการกรม",
        doc_number="ศธ ๐๒๐๑/๑๒๓",
        date=thai_date(datetime.date(2026, 6, 29)),
        subject="ขออนุมัติเบิกค่าใช้จ่ายในการเดินทางไปราชการ",
        to="ผู้อำนวยการกองคลัง",
        body=[
            "ด้วยข้าพเจ้ามีความจำเป็นต้องเดินทางไปราชการเพื่อเข้าร่วมประชุม "
            "ในระหว่างวันที่ ๑–๓ กรกฎาคม ๒๕๖๙ ณ จังหวัดเชียงใหม่",
            "จึงเรียนมาเพื่อโปรดพิจารณาอนุมัติ จะเป็นพระคุณยิ่ง",
        ],
        signer_name="นางสาวเนตรทิพย์ ใจดี",
        signer_position="นักวิชาการเงินและบัญชีชำนาญการ",
        thai_digits=False,
    )
    print(f"✅ สร้างไฟล์ตัวอย่าง: {out}")
